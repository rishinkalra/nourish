#!/usr/bin/env python3
"""Extract Project Nourish requirements into machine-readable and reviewable registers."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path

from docx import Document


REQUIREMENT_ID = re.compile(r"^(?:IOS|ADM|SUB|PRIV|NFR)-[A-Z0-9-]+$")


def cleaned(text: str) -> str:
    return " / ".join(part.strip() for part in text.splitlines() if part.strip())


def rows(table):
    return [[cleaned(cell.text) for cell in row.cells] for row in table.rows]


def dictionaries(table_rows):
    header = table_rows[0]
    return [dict(zip(header, row)) for row in table_rows[1:] if any(row)]


def find_sections(document: Document):
    sections = {
        "consumer_endpoints": [],
        "admin_endpoints": [],
        "error_categories": [],
        "analytics_events": [],
        "decisions": [],
        "notifications": [],
        "data_entities": [],
        "architecture": [],
        "security_controls": [],
        "test_strategy": [],
    }
    requirements = []
    endpoint_tables_seen = 0

    for table_number, table in enumerate(document.tables, 1):
        table_rows = rows(table)
        if not table_rows:
            continue
        header = table_rows[0]

        for row_number, row in enumerate(table_rows[1:], 2):
            if row and REQUIREMENT_ID.match(row[0]):
                if row[0].startswith("NFR-"):
                    requirement = {
                        "id": row[0],
                        "priority": "Release target",
                        "area": row[1] if len(row) > 1 else "",
                        "requirement": row[2] if len(row) > 2 else "",
                    }
                elif len(row) >= 3:
                    requirement = {
                        "id": row[0],
                        "priority": row[1] if len(row) > 1 else "",
                        "area": row[0].split("-")[1],
                        "requirement": row[2] if len(row) > 2 else "",
                    }
                else:
                    requirement = {
                        "id": row[0],
                        "priority": "Release control",
                        "area": row[0].split("-")[1],
                        "requirement": row[1] if len(row) > 1 else "",
                    }
                requirement["source"] = f"Table {table_number}, row {row_number}"
                requirements.append(requirement)

        if header == ["Endpoint", "Purpose"]:
            endpoint_tables_seen += 1
            key = "consumer_endpoints" if endpoint_tables_seen == 1 else "admin_endpoints"
            sections[key] = dictionaries(table_rows)
        elif header == ["Code", "Meaning"]:
            sections["error_categories"] = dictionaries(table_rows)
        elif header == ["Event", "Trigger", "Key properties"]:
            sections["analytics_events"] = dictionaries(table_rows)
        elif header == ["Decision area", "Decision needed"]:
            sections["decisions"] = dictionaries(table_rows)
        elif header == ["Message", "Trigger", "Action"]:
            sections["notifications"] = dictionaries(table_rows)
        elif header == ["Entity", "Purpose"]:
            sections["data_entities"] = dictionaries(table_rows)
        elif header == ["Layer", "Recommendation"]:
            sections["architecture"] = dictionaries(table_rows)
        elif header == ["Area", "Minimum control"]:
            sections["security_controls"] = dictionaries(table_rows)
        elif header == ["Test level", "Coverage"]:
            sections["test_strategy"] = dictionaries(table_rows)

    release_acceptance = []
    collecting = False
    for paragraph in document.paragraphs:
        text = cleaned(paragraph.text)
        if text == "19.2 Release-blocking acceptance criteria":
            collecting = True
            continue
        if text == "19.3 Definition of done for a feature":
            collecting = False
        elif collecting and text:
            release_acceptance.append(text)

    sections["release_acceptance"] = release_acceptance
    return requirements, sections


def md_table(headers, body):
    lines = ["| " + " | ".join(headers) + " |", "|" + "|".join(["---"] * len(headers)) + "|"]
    for row in body:
        lines.append("| " + " | ".join(str(value).replace("|", "\\|") for value in row) + " |")
    return "\n".join(lines)


def render_markdown(register, coverage):
    requirements = register["requirements"]
    status_counts = Counter(item["status"] for item in requirements)
    priority_counts = Counter(item["priority"] for item in requirements)
    lines = [
        "# Project Nourish requirements traceability",
        "",
        f"Source: `{register['source']['name']}` · SHA-256 `{register['source']['sha256'][:12]}…` · extracted {register['generated_at']}",
        "",
        "This register is generated from the product specification. Every numbered requirement defaults to `not_started` unless implementation evidence is explicitly recorded in `docs/coverage_overrides.json`.",
        "",
        "## Coverage summary",
        "",
        f"- Numbered requirements: **{len(requirements)}**",
        f"- Priorities: {', '.join(f'{key}: {value}' for key, value in sorted(priority_counts.items()))}",
        f"- Statuses: {', '.join(f'{key}: {value}' for key, value in sorted(status_counts.items()))}",
        f"- Consumer endpoints: **{len(register['consumer_endpoints'])}**",
        f"- Admin endpoints: **{len(register['admin_endpoints'])}**",
        f"- Analytics events: **{len(register['analytics_events'])}**",
        f"- Release-blocking acceptance criteria: **{len(register['release_acceptance'])}**",
        f"- Unresolved product decisions: **{len(register['decisions'])}**",
        "",
        "## Numbered requirements",
        "",
        md_table(
            ["ID", "Priority", "Requirement / target", "Status", "Evidence", "Notes"],
            [[item["id"], item["priority"], item["requirement"], item["status"], item.get("evidence", "—"), item.get("notes", "")] for item in requirements],
        ),
    ]

    section_specs = [
        ("Consumer API endpoints", "consumer_endpoints", ["Endpoint", "Purpose"]),
        ("Admin API endpoints", "admin_endpoints", ["Endpoint", "Purpose"]),
        ("Structured API errors", "error_categories", ["Code", "Meaning"]),
        ("Analytics event catalogue", "analytics_events", ["Event", "Trigger", "Key properties"]),
        ("Notification catalogue", "notifications", ["Message", "Trigger", "Action"]),
        ("Core data entities", "data_entities", ["Entity", "Purpose"]),
        ("Architecture decisions", "architecture", ["Layer", "Recommendation"]),
        ("Security control baseline", "security_controls", ["Area", "Minimum control"]),
        ("Test strategy", "test_strategy", ["Test level", "Coverage"]),
        ("Decisions required before implementation", "decisions", ["Decision area", "Decision needed"]),
    ]
    for title, key, headers in section_specs:
        lines += ["", f"## {title}", ""]
        values = register[key]
        tracked = key not in {"architecture", "decisions"}
        if tracked:
            overrides = coverage.get("__interfaces__", {}).get(key, {})
            rendered_headers = headers + ["Status", "Evidence"]
            rendered_rows = []
            for item in values:
                override = overrides.get(item.get(headers[0], ""), {})
                rendered_rows.append(
                    [item.get(header, "") for header in headers]
                    + [override.get("status", "not_started"), override.get("evidence", "—")]
                )
            lines.append(md_table(rendered_headers, rendered_rows))
        else:
            lines.append(md_table(headers, [[item.get(header, "") for header in headers] for item in values]))

    lines += ["", "## Release-blocking acceptance criteria", ""]
    lines.extend(f"- [ ] {criterion}" for criterion in register["release_acceptance"])
    lines += [
        "",
        "## Status discipline",
        "",
        "- `verified`: implemented and covered by an executable check or build verification.",
        "- `partial`: some required behavior exists, but the complete requirement is not met.",
        "- `design_only`: represented in prototype/design, without production behavior.",
        "- `contract_only`: typed route, schema, or error contract exists, without production service behavior.",
        "- `not_started`: no implementation evidence recorded.",
        "- `blocked`: cannot proceed until a listed business, legal, nutrition, or vendor decision is resolved.",
    ]
    return "\n".join(lines) + "\n"


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--coverage", type=Path, required=True)
    parser.add_argument("--json-out", type=Path, required=True)
    parser.add_argument("--md-out", type=Path, required=True)
    args = parser.parse_args()

    coverage = json.loads(args.coverage.read_text())
    document = Document(args.docx)
    requirements, sections = find_sections(document)
    for requirement in requirements:
        override = coverage.get(requirement["id"], {})
        requirement.update({"status": override.get("status", "not_started"), "evidence": override.get("evidence", "—"), "notes": override.get("notes", "")})

    register = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "source": {"name": args.docx.name, "sha256": hashlib.sha256(args.docx.read_bytes()).hexdigest()},
        "requirements": requirements,
        **sections,
    }
    args.json_out.parent.mkdir(parents=True, exist_ok=True)
    args.md_out.parent.mkdir(parents=True, exist_ok=True)
    args.json_out.write_text(json.dumps(register, indent=2, ensure_ascii=False) + "\n")
    args.md_out.write_text(render_markdown(register, coverage))


if __name__ == "__main__":
    main()
